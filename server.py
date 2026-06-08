import os
import sys
import subprocess
import threading
import re
import tempfile
import uuid
from pathlib import Path
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS

app = Flask(__name__, static_folder=".")
CORS(app)

SCRIPTS_DIR = Path("generated_scripts")
SCRIPTS_DIR.mkdir(exist_ok=True)

MEDIA_DIR = Path("media")
MEDIA_DIR.mkdir(exist_ok=True)

render_status = {}


# ══════════════════════════════════════════════════════════════
#  BOOKMARK DETECTION
#  Handles all three string formats used by generated code:
#    1. """..."""  triple-quoted
#    2. text='...' single-quoted
#    3. text=(...) parenthesised multi-line  ← primary format
# ══════════════════════════════════════════════════════════════

def get_all_bookmark_tags(code):
    """
    Find all <bookmark mark='X'/> tags inside voiceover text strings.
    Recognises triple-quoted, single-quoted, and parenthesised formats.
    """
    found = set()
    BM_PAT = re.compile(
        r"<bookmark\s+mark=['\"]([^'\"]+)['\"]\s*/?>",
        re.IGNORECASE,
    )

    # ── 1. Triple-quoted strings ──────────────────────────────
    for m in re.finditer(r'"""([\s\S]*?)"""', code):
        for bm in BM_PAT.finditer(m.group(1)):
            found.add(bm.group(1))

    # ── 2. Single-quoted text='...' strings ──────────────────
    for m in re.finditer(r"text\s*=\s*r?'([^']*)'", code):
        for bm in BM_PAT.finditer(m.group(1)):
            found.add(bm.group(1))

    # ── 3. Parenthesised text=(...) multi-line strings ────────
    #    Matches: text=(  'string1 '  'string2 '  )
    #    Captures everything inside the outer parens.
    for m in re.finditer(
        r'text\s*=\s*\(([\s\S]*?)\)',
        code,
    ):
        inner = m.group(1)
        # Only process if it looks like a string (contains quotes)
        if "'" in inner or '"' in inner:
            for bm in BM_PAT.finditer(inner):
                found.add(bm.group(1))

    return found


def get_all_wait_calls(code):
    """Find all wait_until_bookmark('X') calls."""
    return re.findall(
        r'self\.wait_until_bookmark\s*\(\s*["\']([^"\']+)["\']\s*\)',
        code,
    )


# ══════════════════════════════════════════════════════════════
#  BOOKMARK INJECTION
#  Tries three strategies in order:
#    A. Triple-quoted string before the wait call
#    B. Single-quoted text='...' string
#    C. Parenthesised text=(...) multi-line string  ← primary
# ══════════════════════════════════════════════════════════════

def inject_bookmark(code, bm_name):
    """
    Find the correct voiceover block before the wait call
    and inject the bookmark tag.
    Returns (new_code, success_bool)
    """
    BM_PAT = re.compile(
        r"<bookmark\s+mark=['\"]" + re.escape(bm_name) + r"['\"]\s*/?>",
        re.IGNORECASE,
    )

    # Locate the wait call
    wait_pat = re.compile(
        r'self\.wait_until_bookmark\s*\(\s*["\']'
        + re.escape(bm_name)
        + r'["\']\s*\)',
    )
    wait_matches = list(wait_pat.finditer(code))
    if not wait_matches:
        return code, False

    wait_pos    = wait_matches[0].start()
    code_before = code[:wait_pos]

    # ── Strategy A: triple-quoted string ─────────────────────
    triple_matches = list(re.finditer(r'"""([\s\S]*?)"""', code_before))
    if triple_matches:
        last        = triple_matches[-1]
        inner       = last.group(1)
        inner_start = last.start() + 3
        inner_end   = last.end()   - 3

        if BM_PAT.search(inner):
            return code, False   # already present

        stripped  = inner.lstrip('\n')
        leading   = inner[: len(inner) - len(stripped)]
        new_inner = leading + f"<bookmark mark='{bm_name}'/> " + stripped
        new_code  = code[:inner_start] + new_inner + code[inner_end:]
        return new_code, True

    # ── Strategy B: single-quoted text='...' ─────────────────
    single_matches = list(
        re.finditer(r"(text\s*=\s*r?')((?:[^'\\]|\\.)*)'", code_before)
    )
    if single_matches:
        last    = single_matches[-1]
        prefix  = last.group(1)
        inner   = last.group(2)
        start   = last.start()
        end     = last.end()

        if BM_PAT.search(inner):
            return code, False

        new_inner = f"<bookmark mark='{bm_name}'/> " + inner
        new_full  = prefix + new_inner + "'"
        new_code  = code[:start] + new_full + code[end:]
        return new_code, True

    # ── Strategy C: parenthesised text=(...) multi-line ──────
    #    This is the primary format used by all generated scenes.
    #
    #    Pattern explanation:
    #      text\s*=\s*\(          — opening: text=(
    #      ([\s\S]*?)             — inner content (lazy)
    #      \)                     — closing )
    #
    #    We find the LAST such block before the wait call,
    #    then inject the bookmark at the start of the first
    #    string literal inside the parens.

    paren_matches = list(
        re.finditer(r'text\s*=\s*\(([\s\S]*?)\)', code_before)
    )
    if paren_matches:
        last      = paren_matches[-1]
        inner     = last.group(1)
        blk_start = last.start()
        blk_end   = last.end()

        if BM_PAT.search(inner):
            return code, False

        # Find position of first opening quote inside the parens
        first_q = re.search(r"['\"]", inner)
        if first_q:
            # Inject AFTER the opening quote character
            q_offset  = first_q.start() + 1          # +1 to step past the quote
            abs_pos   = blk_start + len("text=(") + q_offset

            # Recalculate carefully using the actual match span
            inner_abs_start = last.start(1)           # start of group(1) in full code
            inject_at       = inner_abs_start + first_q.start() + 1

            new_code = (
                code[:inject_at]
                + f"<bookmark mark='{bm_name}'/> "
                + code[inject_at:]
            )
            return new_code, True

    return code, False


# ══════════════════════════════════════════════════════════════
#  AUTO-REPAIR ENGINE
# ══════════════════════════════════════════════════════════════

def auto_repair_bookmarks(code, filename=''):
    """
    Keep injecting missing bookmarks until all wait calls are matched.
    Last resort: remove unfixable wait calls so Manim never crashes.
    Returns (repaired_code, fixes_list, errors_list)
    """
    fixes  = []
    errors = []

    for _ in range(120):
        existing = get_all_bookmark_tags(code)
        waits    = get_all_wait_calls(code)

        missing = next((w for w in waits if w not in existing), None)
        if missing is None:
            break

        print(f'  [BookmarkFix] Injecting: {missing}')
        new_code, ok = inject_bookmark(code, missing)

        if ok:
            code = new_code
            fixes.append(missing)
        else:
            # Last resort: remove the broken wait call
            rm_pat  = (
                r'[ \t]*self\.wait_until_bookmark\s*\(\s*["\']'
                + re.escape(missing)
                + r'["\']\s*\)\s*\n?'
            )
            cleaned = re.sub(rm_pat, '\n', code)
            if cleaned != code:
                code = cleaned
                msg  = f"Removed unfixable wait_until_bookmark('{missing}')"
                errors.append(msg)
                print(f'  [BookmarkFix] ⚠️  {msg}')
            else:
                errors.append(f"Could not fix or remove '{missing}'")
            break

    # ── Final safety pass ─────────────────────────────────────
    final_existing = get_all_bookmark_tags(code)
    final_waits    = get_all_wait_calls(code)
    still_missing  = [w for w in final_waits if w not in final_existing]

    for bm in still_missing:
        rm_pat = (
            r'[ \t]*self\.wait_until_bookmark\s*\(\s*["\']'
            + re.escape(bm)
            + r'["\']\s*\)\s*\n?'
        )
        code = re.sub(rm_pat, '\n', code)
        msg  = f"Final pass: removed unfixable wait_until_bookmark('{bm}')"
        errors.append(msg)
        print(f'  [BookmarkFix] 🔴 {msg}')

    return code, fixes, errors


# ══════════════════════════════════════════════════════════════
#  WEIGHT= FIX
# ══════════════════════════════════════════════════════════════

def fix_weight_argument(code):
    """Remove weight= keyword from non-Text Mobjects."""
    fixes     = []
    new_lines = []

    for i, line in enumerate(code.split('\n')):
        # Text() and MarkupText() are allowed to keep weight=
        if re.search(r'\b(?:MarkupText|Text)\s*\(', line):
            new_lines.append(line)
            continue

        if re.search(r'\bweight\s*=', line):
            fixed = line
            fixed = re.sub(
                r',\s*weight\s*=\s*(?:"[^"]*"|\'[^\']*\'|[A-Z_][A-Z_0-9]*|\d+)',
                '', fixed,
            )
            fixed = re.sub(
                r'\bweight\s*=\s*(?:"[^"]*"|\'[^\']*\'|[A-Z_][A-Z_0-9]*|\d+)\s*,\s*',
                '', fixed,
            )
            fixed = re.sub(
                r'\bweight\s*=\s*(?:"[^"]*"|\'[^\']*\'|[A-Z_][A-Z_0-9]*|\d+)',
                '', fixed,
            )
            if fixed != line:
                fixes.append(f'Line {i + 1}: removed weight=')
            new_lines.append(fixed)
        else:
            new_lines.append(line)

    return '\n'.join(new_lines), fixes


# ══════════════════════════════════════════════════════════════
#  MASTER PREPARE FUNCTION
#  Called before every save AND every render.
# ══════════════════════════════════════════════════════════════

def prepare_script(code, filename=''):
    """
    Run ALL auto-fixes.
    Returns (fixed_code, report_dict)
    """
    report = {
        'bookmark_fixes' : [],
        'bookmark_errors': [],
        'weight_fixes'   : [],
        'waits_before'   : 0,
        'tags_before'    : 0,
        'waits_after'    : 0,
        'tags_after'     : 0,
    }

    report['waits_before'] = len(get_all_wait_calls(code))
    report['tags_before']  = len(get_all_bookmark_tags(code))

    # Fix 1 — weight= on non-Text Mobjects
    code, wfixes = fix_weight_argument(code)
    report['weight_fixes'] = wfixes

    # Fix 2 — missing bookmark tags
    code, bfixes, berrors = auto_repair_bookmarks(code, filename)
    report['bookmark_fixes']  = bfixes
    report['bookmark_errors'] = berrors

    report['waits_after'] = len(get_all_wait_calls(code))
    report['tags_after']  = len(get_all_bookmark_tags(code))

    return code, report


# ══════════════════════════════════════════════════════════════
#  TEXT NORMALISATION
# ══════════════════════════════════════════════════════════════

def normalize_text(raw_text):
    text = (raw_text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ══════════════════════════════════════════════════════════════
#  TRANSCRIPT PARSER
#  Fixed: Summary block is NO LONGER stripped — it is part of
#  the transcript that Claude needs to generate the animation.
# ══════════════════════════════════════════════════════════════

def parse_transcripts_from_text(raw_text):
    text = normalize_text(raw_text)
    if not text:
        return []

    transcripts = []

    METADATA_FIELDS = (
        r'Subject|Grade|Topic|Subtopic|'
        r'Learning\s*Outcomes?|Outcomes?|'
        r'Estimated\s*Duration|Duration'
    )

    def is_metadata_label(line):
        t = (line or "").strip()
        return bool(re.match(
            r'^(' + METADATA_FIELDS + r')\s*[:\t]?\s*$',
            t, re.IGNORECASE,
        ))

    def is_subject_start(line):
        t = (line or "").strip()
        if not t:
            return False
        if re.match(r'^Subject\s*[:\t]?\s*$',  t, re.IGNORECASE): return True
        if re.match(r'^Subject\s*[:\t]\s*.+$', t, re.IGNORECASE): return True
        if re.match(r'^Subject\s+.+$',          t, re.IGNORECASE): return True
        return False

    def extract_field(lines, field_name):
        rx = re.compile(
            r'^' + field_name + r'\s*[:\t]?\s*(.*)$',
            re.IGNORECASE,
        )
        for i, line in enumerate(lines):
            m = rx.match(line.strip())
            if not m:
                continue
            val = (m.group(1) or "").strip()
            if val:
                return val
            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if nxt and not is_metadata_label(nxt):
                    return nxt
        return ""

    def mark_consumed(lines, field_regex, consumed):
        rx_label = re.compile(
            r'^' + field_regex + r'\s*[:\t]?\s*$', re.IGNORECASE)
        rx_sep   = re.compile(
            r'^' + field_regex + r'\s*[:\t]\s*.+$', re.IGNORECASE)
        rx_space = re.compile(
            r'^' + field_regex + r'\s+.+$',          re.IGNORECASE)
        for i, line in enumerate(lines):
            t = line.strip()
            if rx_label.match(t):
                consumed.add(i)
                if i + 1 < len(lines):
                    nxt = lines[i + 1].strip()
                    if nxt and not is_metadata_label(nxt):
                        consumed.add(i + 1)
            elif rx_sep.match(t) or rx_space.match(t):
                consumed.add(i)

    # ── Split raw text into per-transcript blocks ─────────────
    all_lines = [ln.rstrip() for ln in text.split("\n")]
    blocks    = []
    current   = []

    for line in all_lines:
        if is_subject_start(line) and current:
            blocks.append("\n".join(current).strip())
            current = [line]
        else:
            current.append(line)
    if current:
        blocks.append("\n".join(current).strip())

    valid_blocks = [
        b.strip() for b in blocks
        if b.strip() and is_subject_start(
            next(
                (ln.strip() for ln in b.split("\n") if ln.strip()),
                "",
            )
        )
    ]

    # ── Parse each block ──────────────────────────────────────
    for idx, block in enumerate(valid_blocks):
        lines = [ln.strip() for ln in block.split("\n") if ln.strip()]

        subject  = extract_field(lines, r'Subject')
        grade    = extract_field(lines, r'Grade')
        topic    = extract_field(lines, r'Topic')
        subtopic = extract_field(lines, r'Subtopic')
        outcomes = (
            extract_field(lines, r'Learning\s*Outcomes?')
            or extract_field(lines, r'Outcomes?')
        )
        duration_raw = (
            extract_field(lines, r'Estimated\s*Duration')
            or extract_field(lines, r'Duration')
        )

        # Mark which lines are metadata so we can exclude them
        consumed = set()
        for field in [
            r'Subject', r'Grade', r'Topic', r'Subtopic',
            r'Learning\s*Outcomes?', r'Outcomes?',
            r'Estimated\s*Duration', r'Duration',
        ]:
            mark_consumed(lines, field, consumed)

        # ── Transcript body = everything NOT in metadata ───────
        # NOTE: Summary block is intentionally KEPT — Claude needs it
        #       to generate the animation's summary segment.
        transcript_body = "\n".join(
            line for i, line in enumerate(lines)
            if i not in consumed
        ).strip()

        if not transcript_body or len(transcript_body) < 30:
            continue

        # ── Parse grade number ─────────────────────────────────
        grade_num   = re.search(r'\d+', grade or '')
        grade_clean = (
            f"Grade {grade_num.group()}"
            if grade_num
            else (grade or "Grade 7")
        )

        # ── Parse duration ─────────────────────────────────────
        dur_secs = 90
        mm = re.search(r'(\d+):(\d+)', duration_raw or '')
        if mm:
            dur_secs = int(mm.group(1)) * 60 + int(mm.group(2))
        else:
            dm = re.search(r'(\d+)', duration_raw or '')
            if dm:
                dur_secs = int(dm.group(1))

        # ── Build title ────────────────────────────────────────
        title = topic or subtopic or subject or f"Transcript {idx + 1}"
        if subtopic and topic and subtopic.lower() != topic.lower():
            title = f"{topic} - {subtopic}"

        transcripts.append({
            "title"      : title,
            "subject"    : subject or "Mathematics",
            "grade"      : grade_clean,
            "topic"      : topic or "",
            "subtopic"   : subtopic or "",
            "outcomes"   : outcomes or "",
            "duration"   : str(dur_secs),
            "transcript" : transcript_body,
            "raw_block"  : block,
        })

    return transcripts


# ══════════════════════════════════════════════════════════════
#  DOCUMENT EXTRACTORS
# ══════════════════════════════════════════════════════════════

def extract_docx_text(path):
    import docx as python_docx
    doc   = python_docx.Document(path)
    parts = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_lines = [
                    p.text.strip()
                    for p in cell.paragraphs
                    if p.text.strip()
                ]
                cell_text = "\n".join(cell_lines).strip()
                row_cells.append(cell_text)
            row_cells = [c for c in row_cells if c]
            if row_cells:
                parts.append("\t".join(row_cells))
    return normalize_text("\n".join(parts))


def extract_pdf_text_from_pymupdf(data):
    import fitz
    doc   = fitz.open(stream=data, filetype="pdf")
    parts = []
    try:
        for page in doc:
            txt = page.get_text("text")
            if txt and txt.strip():
                parts.append(txt)
    finally:
        doc.close()
    return normalize_text("\n".join(parts))


# ══════════════════════════════════════════════════════════════
#  FLASK ROUTES
# ══════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return send_from_directory(".", "index.html")


@app.route("/media/<path:filepath>")
def serve_media(filepath):
    return send_from_directory(str(MEDIA_DIR.resolve()), filepath)


@app.route("/api/health")
def health():
    try:
        result = subprocess.run(
            ["manim", "--version"],
            capture_output=True, text=True, timeout=10,
        )
        manim_version = result.stdout.strip() or result.stderr.strip()
        manim_ok      = result.returncode == 0
    except FileNotFoundError:
        manim_version = "NOT FOUND"
        manim_ok      = False
    except Exception as e:
        manim_version = str(e)
        manim_ok      = False

    docx_ok = False
    pdf_ok  = False
    try:
        import docx  # noqa
        docx_ok = True
    except ImportError:
        pass
    try:
        import fitz  # noqa
        pdf_ok = True
    except ImportError:
        pass

    return jsonify({
        "status"        : "ok",
        "manim_ok"      : manim_ok,
        "manim_version" : manim_version,
        "scripts_dir"   : str(SCRIPTS_DIR.resolve()),
        "scripts_count" : len(list(SCRIPTS_DIR.glob("*.py"))),
        "python"        : sys.executable,
        "docx_ok"       : docx_ok,
        "pdf_ok"        : pdf_ok,
    })


@app.route("/api/parse-document", methods=["POST"])
def parse_document():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f        = request.files["file"]
    filename = (f.filename or "").lower()
    raw_text = ""

    try:
        if filename.endswith(".txt") or filename.endswith(".text"):
            raw_text = f.read().decode("utf-8", errors="replace")

        elif filename.endswith(".docx"):
            try:
                tmp = tempfile.NamedTemporaryFile(
                    delete=False, suffix=".docx")
                tmp.write(f.read())
                tmp.close()
                try:
                    raw_text = extract_docx_text(tmp.name)
                finally:
                    try:
                        os.unlink(tmp.name)
                    except OSError:
                        pass
            except ImportError:
                return jsonify({
                    "error": (
                        "python-docx not installed. "
                        "Run: pip install python-docx"
                    )
                }), 500

        elif filename.endswith(".pdf"):
            try:
                raw_text = extract_pdf_text_from_pymupdf(f.read())
            except ImportError:
                return jsonify({
                    "error": (
                        "PyMuPDF not installed. "
                        "Run: pip install PyMuPDF"
                    )
                }), 500

        else:
            return jsonify({
                "error": "Unsupported file type. Use .txt .docx or .pdf"
            }), 400

    except Exception as e:
        return jsonify({"error": f"File read error: {str(e)}"}), 500

    raw_text = normalize_text(raw_text)
    if not raw_text.strip():
        return jsonify({"error": "Document appears empty"}), 400

    transcripts = parse_transcripts_from_text(raw_text)
    if not transcripts:
        return jsonify({
            "error"  : (
                "No transcripts found. "
                "Each transcript must start with 'Subject'."
            ),
            "preview": raw_text[:1200],
        }), 400

    return jsonify({
        "success"    : True,
        "count"      : len(transcripts),
        "transcripts": transcripts,
    })


@app.route("/api/save-script", methods=["POST"])
def save_script():
    data     = request.json or {}
    code     = data.get("code", "")
    filename = data.get("filename", "coschool_script.py")

    filename = filename.replace("/", "").replace("\\", "")
    if not filename.endswith(".py"):
        filename += ".py"

    # Run auto-fixes when saving
    fixed_code, report = prepare_script(code, filename)

    total_fixes = (
        len(report['bookmark_fixes']) + len(report['weight_fixes'])
    )
    if total_fixes > 0:
        print(f'[Save] Auto-fixed {total_fixes} issue(s) in {filename}')
        for fix in report['bookmark_fixes']:
            print(f'  ✅ Bookmark injected: {fix}')
        for fix in report['weight_fixes']:
            print(f'  ✅ Weight= removed: {fix}')

    path = SCRIPTS_DIR / filename
    path.write_text(fixed_code, encoding="utf-8")

    return jsonify({
        "success" : True,
        "filename": filename,
        "path"    : str(path.resolve()),
        "fixes"   : report,
    })


@app.route("/api/get-script/<filename>")
def get_script(filename):
    filename = os.path.basename(filename)
    path     = SCRIPTS_DIR / filename
    if not path.exists():
        return jsonify({"success": False, "error": "Not found"}), 404
    code = path.read_text(encoding="utf-8")
    return jsonify({"success": True, "code": code, "filename": filename})


@app.route("/api/list-scripts")
def list_scripts():
    scripts = []
    for f in sorted(
        SCRIPTS_DIR.glob("*.py"),
        key=lambda x: x.stat().st_mtime,
        reverse=True,
    ):
        scripts.append({
            "filename": f.name,
            "size_kb" : round(f.stat().st_size / 1024, 1),
            "path"    : str(f.resolve()),
        })
    return jsonify({"scripts": scripts})


@app.route("/api/render", methods=["POST"])
def render():
    data      = request.json or {}
    filename  = data.get("filename", "")
    classname = data.get("classname", "")
    quality   = data.get("quality", "m")
    job_id    = data.get("job_id", str(uuid.uuid4()))

    flags = {"l": "-ql", "m": "-qm", "h": "-qh", "k": "-qk"}
    flag  = flags.get(quality, "-qm")

    script_path = SCRIPTS_DIR / filename
    if not script_path.exists():
        return jsonify({"error": f"Script not found: {filename}"}), 404

    render_status[job_id] = {
        "state"    : "running",
        "log"      : [],
        "percent"  : 0,
        "output"   : "",
        "video_url": "",
        "error"    : "",
        "fixes"    : {},
    }

    def do_render():
        try:
            # ── Step 1: Read script ───────────────────────────
            original_code = script_path.read_text(encoding="utf-8")

            # ── Step 2: Auto-fix ──────────────────────────────
            print(f'\n[Render] Auto-fixing: {filename}')
            render_status[job_id]['log'].append(
                f'[AutoFix] Scanning {filename} for issues…'
            )

            fixed_code, report = prepare_script(original_code, filename)
            render_status[job_id]['fixes'] = report

            # Diagnostic counts
            render_status[job_id]['log'].append(
                f'[AutoFix] waits={report["waits_before"]} '
                f'tags_before={report["tags_before"]} '
                f'tags_after={report["tags_after"]}'
            )

            total_fixes = (
                len(report['bookmark_fixes']) +
                len(report['weight_fixes'])
            )

            if total_fixes > 0:
                script_path.write_text(fixed_code, encoding="utf-8")
                for fix in report['bookmark_fixes']:
                    msg = f'[AutoFix] ✅ Bookmark injected: {fix}'
                    render_status[job_id]['log'].append(msg)
                    print(f'  {msg}')
                for fix in report['weight_fixes']:
                    msg = f'[AutoFix] ✅ Weight= removed: {fix}'
                    render_status[job_id]['log'].append(msg)
                    print(f'  {msg}')
                render_status[job_id]['log'].append(
                    f'[AutoFix] Fixed {total_fixes} issue(s) — saved'
                )
                print(f'[Render] Fixed {total_fixes} issue(s), saved.')
            else:
                render_status[job_id]['log'].append(
                    '[AutoFix] ✅ No issues found — script is clean'
                )
                print('[Render] No issues found — script is clean')

            for err in report['bookmark_errors']:
                msg = f'[AutoFix] ⚠️  {err}'
                render_status[job_id]['log'].append(msg)
                print(f'  {msg}')

            # ── Step 3: Manim render ──────────────────────────
            cmd = [
                sys.executable, "-m", "manim",
                "render", flag,
                str(script_path.resolve()),
                classname,
            ]

            render_status[job_id]['log'].append(
                f'[Manim] Running: {" ".join(cmd)}'
            )
            print(f'[Render] CMD: {" ".join(cmd)}')

            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                cwd=str(Path(".").resolve()),
            )

            if process.stdout is not None:
                for line in process.stdout:
                    line = line.rstrip()
                    render_status[job_id]['log'].append(line)

                    m = re.search(r'(\d+)%', line)
                    if m:
                        render_status[job_id]['percent'] = int(m.group(1))

                    if "File ready at" in line or "Rendered" in line:
                        mp4 = re.search(r"'([^']+\.mp4)'", line)
                        if not mp4:
                            mp4 = re.search(r'"([^"]+\.mp4)"', line)
                        if mp4:
                            render_status[job_id]['output'] = mp4.group(1)

            process.wait()

            # ── Step 4: Result ────────────────────────────────
            if process.returncode == 0:
                render_status[job_id]['state']   = "done"
                render_status[job_id]['percent'] = 100

                res_map = {
                    "l": "480p15",
                    "m": "720p30",
                    "h": "1080p60",
                    "k": "2160p60",
                }
                res      = res_map.get(quality, "720p30")
                base     = filename.replace(".py", "")
                expected = (
                    MEDIA_DIR / "videos" / base / res / f"{classname}.mp4"
                )

                if expected.exists():
                    render_status[job_id]['output']    = str(expected.resolve())
                    render_status[job_id]['video_url'] = (
                        f"videos/{base}/{res}/{classname}.mp4"
                    )
                else:
                    found = list(MEDIA_DIR.rglob(f"{classname}.mp4"))
                    if found:
                        mp4_path = found[0]
                        render_status[job_id]['output']    = str(mp4_path.resolve())
                        render_status[job_id]['video_url'] = str(
                            mp4_path.relative_to(MEDIA_DIR)
                        ).replace("\\", "/")

                print(f'[Render] ✅ Done: {filename}')

            else:
                render_status[job_id]['state'] = "error"
                render_status[job_id]['error'] = "Render failed. Check log."
                print(f'[Render] ❌ Failed: {filename}')

        except Exception as e:
            render_status[job_id]['state'] = "error"
            render_status[job_id]['error'] = str(e)
            print(f'[Render] ❌ Exception: {e}')

    threading.Thread(target=do_render, daemon=True).start()
    return jsonify({"success": True, "job_id": job_id})


@app.route("/api/render-status/<job_id>")
def get_render_status(job_id):
    return jsonify(render_status.get(job_id, {
        "state"    : "unknown",
        "log"      : [],
        "percent"  : 0,
        "video_url": "",
    }))


# ══════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("  Coschool Pipeline Studio Server v14")
    print("  http://127.0.0.1:5000")
    print("=" * 60)
    print(f"  Scripts : {SCRIPTS_DIR.resolve()}")
    print(f"  Media   : {MEDIA_DIR.resolve()}")
    print(f"  Python  : {sys.executable}")
    print("=" * 60)
    print("  AUTO-FIX enabled on every save + render:")
    print("    ✅ Bookmark tags — all 3 string formats detected")
    print("       • triple-quoted  \"\"\"...\"\"\"")
    print("       • single-quoted  text='...'")
    print("       • parenthesised  text=(...)  ← primary format")
    print("    ✅ weight= errors removed automatically")
    print("    ✅ Summary block preserved in transcript body")
    print("=" * 60)
    print("  Install:")
    print("    pip install flask flask-cors python-docx PyMuPDF")
    print("=" * 60)
    app.run(debug=False, port=5000, host="127.0.0.1")